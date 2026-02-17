"""
Circular views – generate body, preview, PDF & DOCX download.
"""
import logging
import os
from io import BytesIO

from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML

from ..data.constants import DESIGNATION_MAP
from ..models import DocumentLog
from ..services.ai_service import get_gemini_model
from ..services.data_loader import get_circular_data
from ..utils_new.formatters import format_date_ddmmyyyy, safe_designation

logger = logging.getLogger('generator')


# -------- AI body generation --------
def generate_circular_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("body_prompt", "").strip()
    lang = request.POST.get("language", "en").strip() or "en"

    if lang == "hi":
        system_prompt = """
आप BISAG-N के लिए एक सरकारी परिपत्र (Circular) का केवल मुख्य भाग (BODY) लिख रहे हैं।

महत्वपूर्ण नियम:
- केवल परिपत्र का मुख्य विषय-वस्तु लिखें।
- कोई विषय (Subject) न लिखें।
- कोई शीर्षक न लिखें।
- कोई संदर्भ संख्या न लिखें।
- कोई हस्ताक्षर न लिखें।
- कोई दिनांक न लिखें।
- कोई "प्रेषक" या "प्राप्तकर्ता" न लिखें।
- 1–2 औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- केवल सादा पाठ में उत्तर दें।
"""
    else:
        system_prompt = """
You are drafting ONLY the BODY content of an official Government Circular for BISAG-N.

IMPORTANT Rules:
- Write ONLY the main body content of the circular.
- Do NOT include any subject line.
- Do NOT include any title or heading.
- Do NOT include reference number.
- Do NOT include signature.
- Do NOT include date.
- Do NOT include From or To sections.
- Write 1–2 formal paragraphs only.
- Official government tone.
- Plain text only.
"""

    try:
        model = get_gemini_model()
        res = model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
        return HttpResponse(res.text.strip())
    except Exception as exc:
        logger.exception("Gemini API error (circular): %s", exc)
        return JsonResponse({"error": "AI generation failed. Please try again."}, status=500)


# -------- Preview --------
def result_circular(request):
    if request.method != "POST":
        return redirect("home")

    CIRCULAR = get_circular_data()
    lang = request.POST.get("language", "en").strip() or "en"
    raw_date = request.POST.get("date")
    formatted_date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")
    subject = request.POST.get("subject")
    body = request.POST.get("body")

    from_position = request.POST.get("from_position")
    from_designation = safe_designation(from_position, lang)

    to_ids = request.POST.getlist("to[]")
    people = CIRCULAR.get("people", [])
    to_people = [p for p in people if str(p["id"]) in to_ids]

    if lang == "hi":
        header = {
            "org_name": CIRCULAR.get("header", {}).get("hindi", {}).get("org_name", ""),
            "ministry": CIRCULAR.get("header", {}).get("hindi", {}).get("ministry", ""),
            "government": CIRCULAR.get("header", {}).get("hindi", {}).get("government", ""),
        }
    else:
        header = {
            "org_name": CIRCULAR.get("header", {}).get("english", {}).get("org_name", ""),
            "ministry": CIRCULAR.get("header", {}).get("english", {}).get("ministry", ""),
            "government": CIRCULAR.get("header", {}).get("english", {}).get("government", ""),
        }

    data = {
        "language": lang,
        "header": header,
        "date": formatted_date,
        "subject": subject,
        "body": body,
        "from": from_designation,
        "to_people": to_people,
    }

    # Log to database
    try:
        DocumentLog.objects.create(
            document_type="Circular",
            language=lang,
            reference_id=f"CIRCULAR-{formatted_date}",
            content=body or "",
        )
    except Exception as exc:
        logger.warning("Failed to log Circular: %s", exc)

    request.session["circular_data"] = data
    return render(request, "generator/result_circular.html", data)


# -------- PDF --------
def download_circular_pdf(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        data["logo_path"] = f"file:///{logo_path.replace(chr(92), '/')}"

    html = render_to_string("generator/pdf_circular.html", data)

    try:
        pdf = HTML(
            string=html,
            base_url=str(settings.BASE_DIR),
        ).write_pdf(
            optimize_images=True,
            jpeg_quality=85,
            presentational_hints=True,
        )
    except Exception as exc:
        logger.exception("PDF generation failed (circular): %s", exc)
        return HttpResponse("PDF generation failed", status=500)

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Circular.pdf"'
    return response


# -------- DOCX --------
def download_circular_docx(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo
    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()

    # Header lines
    for line in data.get("header", {}).values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Circular title
    lang = data.get("language", "en")
    title_text = "परिपत्र" if lang == "hi" else "Circular"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)

    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data.get('subject', '')}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data.get("body", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From section
    p = doc.add_paragraph(data.get("from", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # To section — Table
    to_people = data.get("to_people", [])
    if to_people:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        sr_label = "क्र." if lang == "hi" else "Sr. No."
        name_label = "नाम" if lang == "hi" else "Name"
        sign_label = "हस्ताक्षर" if lang == "hi" else "Sign"

        hdr_cells[0].text = sr_label
        hdr_cells[1].text = name_label
        hdr_cells[2].text = sign_label

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, person in enumerate(to_people, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            name = person.get("name_hi") if lang == "hi" else person.get("name_en")
            row_cells[1].text = name or ""
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[2].text = ""
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="Circular.docx"'
    return response
