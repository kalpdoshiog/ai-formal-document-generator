"""
Office Order views – generate body, preview, PDF & DOCX download.
"""
import logging
from datetime import date
from io import BytesIO

from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.template.loader import render_to_string

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS

from ..data.constants import DESIGNATION_MAP
from ..models import DocumentLog
from ..services.ai_service import get_gemini_model
from ..services.data_loader import get_office_order_data
from ..utils_new.formatters import format_date_ddmmyyyy, safe_designation

logger = logging.getLogger('generator')


# -------- AI body generation --------
def generate_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("body_prompt", "").strip()
    lang = request.POST.get("language", "en").strip() or "en"
    today = date.today().strftime("%d-%m-%Y")

    if lang == "hi":
        system_prompt = f"""
आप BISAG-N के लिए एक आधिकारिक कार्यालय आदेश की मुख्य सामग्री लिख रहे हैं।

नियम:
- कम से कम 2–3 वाक्यों का एक औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- कोई शीर्षक, संदर्भ, दिनांक, प्रेषक या प्राप्तकर्ता न लिखें।
- बुलेट या क्रमांक का प्रयोग न करें।
- केवल सादा पाठ में उत्तर दें।

Today's Date: {today}
"""
    else:
        system_prompt = f"""
You are drafting the BODY of an official government Office Order for BISAG-N.

Rules:
- Write one formal paragraph (minimum 2–3 sentences).
- Use official government tone.
- Do not include title, reference, date, From or To.
- No bullet points or numbering.
- Plain text only.

Today's Date: {today}
"""

    full_prompt = system_prompt + "\n\nTopic:\n" + prompt

    try:
        model = get_gemini_model()
        res = model.generate_content(full_prompt)
        return HttpResponse(res.text.strip())
    except Exception as exc:
        logger.exception("Gemini API error (office order): %s", exc)
        return JsonResponse({"error": "AI generation failed. Please try again."}, status=500)


# -------- Preview --------
def result_office_order(request):
    if request.method != "POST":
        return redirect("home")

    OFFICE_ORDER = get_office_order_data()
    lang = request.POST.get("language", "en").strip() or "en"
    raw_date = request.POST.get("date")
    formatted_date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")

    reference = request.POST.get("reference", "").strip()
    if not reference:
        reference = (
            "बायसेग-एन/कार्यालय आदेश/2026/"
            if lang == "hi"
            else "BISAG-N/Office Order/2026/"
        )

    data = {
        "language": lang,
        "header": OFFICE_ORDER.get("header", {}).get(lang, []),
        "title": OFFICE_ORDER.get("title_hi") if lang == "hi" else OFFICE_ORDER.get("title_en"),
        "reference": reference,
        "date": formatted_date,
        "body": request.POST.get("body", "").strip(),
        "from": safe_designation(request.POST.get("from_position"), lang),
        "to": [safe_designation(x, lang) for x in request.POST.getlist("to_recipients[]")],
    }

    # Log to database
    try:
        DocumentLog.objects.create(
            document_type="Office Order",
            language=lang,
            reference_id=reference,
            content=data["body"],
        )
    except Exception as exc:
        logger.warning("Failed to log Office Order: %s", exc)

    request.session["doc_data"] = data
    return render(request, "generator/result_office_order.html", data)


# -------- PDF --------
def download_pdf(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    html = render_to_string("generator/pdf_office_order.html", data)

    try:
        pdf = HTML(
            string=html,
            base_url=str(settings.BASE_DIR),
        ).write_pdf(
            stylesheets=[
                CSS(string="""
                    @page { size: A4; margin: 2.5cm; }
                    body { font-family: serif; font-size: 12pt; line-height: 1.6; }
                    .center { text-align: center; }
                    .bold { font-weight: bold; }
                    .ref-date-row { display: table; width: 100%; margin: 20px 0; }
                    .ref-left { display: table-cell; text-align: left; font-weight: bold; width: 50%; }
                    .date-right { display: table-cell; text-align: right; font-weight: bold; width: 50%; }
                    .title { text-align: center; font-weight: bold; text-decoration: underline; margin: 20px 0; }
                    .body { text-align: justify; margin: 20px 0; }
                    .from-section { text-align: right; font-weight: bold; margin: 40px 0 20px; }
                    .to-section { margin-top: 20px; }
                    .to-section div { margin: 5px 0; }
                """)
            ]
        )
    except Exception as exc:
        logger.exception("PDF generation failed (office order): %s", exc)
        return HttpResponse("PDF generation failed", status=500)

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Office_Order.pdf"'
    return response


# -------- DOCX --------
def download_docx(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    doc = Document()

    # Header
    for line in data.get("header", []):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Reference & Date
    p = doc.add_paragraph(f"Ref: {data['reference']}")
    p.runs[0].bold = True

    p = doc.add_paragraph(f"Date: {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    # Title
    p = doc.add_paragraph(data.get("title", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True

    # Body
    doc.add_paragraph(data.get("body", ""))

    # From
    p = doc.add_paragraph(data.get("from", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    # To
    for t in data.get("to", []):
        p = doc.add_paragraph(t)
        p.runs[0].bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="Office_Order.docx"'
        },
    )
