"""
Policy views – generate body, preview, PDF & DOCX download.
"""
import logging
import os
import uuid
import base64
from io import BytesIO

from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML
from PyPDF2 import PdfReader, PdfWriter

from ..data.constants import DESIGNATION_MAP
from ..models import DocumentLog
from ..services.ai_service import get_gemini_model
from ..services.data_loader import get_policy_data
from ..utils_new.formatters import format_date_ddmmyyyy, safe_designation

logger = logging.getLogger('generator')


# -------- AI body generation --------
def generate_policy_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("body_prompt", "").strip()
    lang = request.POST.get("language", "en")

    if lang == "hi":
        system_prompt = """
आप BISAG-N के लिए एक सरकारी नीति (Policy) का केवल मुख्य भाग (BODY) लिख रहे हैं।

महत्वपूर्ण नियम:
- केवल नीति का मुख्य विषय-वस्तु लिखें।
- कोई विषय (Subject) न लिखें।
- कोई शीर्षक न लिखें।
- कोई संदर्भ संख्या न लिखें।
- कोई हस्ताक्षर न लिखें।
- कोई दिनांक न लिखें।
- कोई "प्रेषक" या "प्राप्तकर्ता" न लिखें।
- 1–2 औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- केवल सादा पाठ में उत्तर दें।
"""
    else:
        system_prompt = """
You are drafting ONLY the BODY content of an official Government Policy for BISAG-N.

IMPORTANT Rules:
- Write ONLY the main body content of the policy.
- Do NOT include any subject line.
- Do NOT include any title or heading.
- Do NOT include reference number.
- Do NOT include signature.
- Do NOT include date.
- Do NOT include From or To sections.
- Write 1–2 formal paragraphs only.
- Official government tone.
- Plain text only.
"""

    try:
        model = get_gemini_model()
        res = model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
        return HttpResponse(res.text.strip())
    except Exception as exc:
        logger.exception("Gemini API error (policy): %s", exc)
        return JsonResponse({"error": "AI generation failed. Please try again."}, status=500)


# -------- Preview --------
def result_policy(request):
    if request.method != "POST":
        return redirect("home")

    POLICY = get_policy_data()
    lang = request.POST.get("language")
    raw_date = request.POST.get("date")
    formatted_date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")
    subject = request.POST.get("subject")
    body = request.POST.get("body")
    attached_pdf_name = request.POST.get("attached_pdf_name", "")

    from_position = request.POST.get("from_position")
    from_designation = safe_designation(from_position, lang)

    to_recipients = request.POST.getlist("to_recipients[]")
    to_designations = [safe_designation(x, lang) for x in to_recipients]

    # Get header based on language
    if lang == "hi":
        header_list = POLICY.get("header", {}).get("hi", [])
    else:
        header_list = POLICY.get("header", {}).get("en", [])

    header = {
        "org_name": header_list[0] if len(header_list) > 0 else "",
        "ministry": header_list[1] if len(header_list) > 1 else "",
        "government": header_list[2] if len(header_list) > 2 else "",
    }

    # Handle PDF upload
    uploaded_pdf = request.FILES.get("policy_pdf")
    pdf_path = None
    pdf_base64 = None
    if uploaded_pdf:
        upload_dir = os.path.join(settings.MEDIA_ROOT, "policy_uploads")
        os.makedirs(upload_dir, exist_ok=True)
        pdf_filename = f"policy_{uuid.uuid4().hex}.pdf"
        pdf_path = os.path.join(upload_dir, pdf_filename)
        with open(pdf_path, "wb+") as destination:
            for chunk in uploaded_pdf.chunks():
                destination.write(chunk)

        with open(pdf_path, "rb") as f:
            pdf_base64 = base64.b64encode(f.read()).decode('utf-8')

    data = {
        "language": lang,
        "header": header,
        "date": formatted_date,
        "subject": subject,
        "body": body,
        "from": from_designation,
        "to_designations": to_designations,
        "attached_pdf_name": attached_pdf_name,
        "uploaded_pdf_path": pdf_path,
    }

    # Log to database
    try:
        DocumentLog.objects.create(
            document_type="Policy",
            language=lang,
            reference_id=f"POLICY-{formatted_date}",
            content=body or "",
        )
    except Exception as exc:
        logger.warning("Failed to log Policy: %s", exc)

    # Store in session (without large base64)
    request.session["policy_data"] = data

    # Add pdf_base64 only for template rendering
    if pdf_base64:
        data["pdf_base64"] = pdf_base64

    return render(request, "generator/result_policy.html", data)


# -------- PDF --------
def download_policy_pdf(request):
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        data["logo_path"] = f"file:///{logo_path.replace(chr(92), '/')}"

    html = render_to_string("generator/pdf_policy.html", data)

    try:
        first_page_pdf = HTML(
            string=html,
            base_url=str(settings.BASE_DIR),
        ).write_pdf(
            optimize_images=True,
            jpeg_quality=85,
            presentational_hints=True,
        )
    except Exception as exc:
        logger.exception("PDF generation failed (policy): %s", exc)
        return HttpResponse("PDF generation failed", status=500)

    # Merge with uploaded PDF if present
    uploaded_pdf_path = data.get("uploaded_pdf_path")

    if uploaded_pdf_path and os.path.exists(uploaded_pdf_path):
        try:
            pdf_writer = PdfWriter()

            first_page_reader = PdfReader(BytesIO(first_page_pdf))
            for page in first_page_reader.pages:
                pdf_writer.add_page(page)

            uploaded_reader = PdfReader(uploaded_pdf_path)
            for page in uploaded_reader.pages:
                pdf_writer.add_page(page)

            output = BytesIO()
            pdf_writer.write(output)
            output.seek(0)
            final_pdf = output.read()

            try:
                os.remove(uploaded_pdf_path)
            except OSError as rm_err:
                logger.warning("Could not remove uploaded PDF: %s", rm_err)
        except Exception as exc:
            logger.exception("Error merging PDFs: %s", exc)
            final_pdf = first_page_pdf
    else:
        final_pdf = first_page_pdf

    response = HttpResponse(final_pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Policy.pdf"'
    return response


# -------- DOCX --------
def download_policy_docx(request):
    data = request.session.get("policy_data")
    if not data:
        return HttpResponse("No policy generated", status=400)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo
    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()

    # Header lines
    for line in data.get("header", {}).values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Policy title
    lang = data.get("language", "en")
    title_text = "नीति" if lang == "hi" else "Policy"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)

    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data.get('subject', '')}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data.get("body", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From section
    p = doc.add_paragraph(data.get("from", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # To section – Bullet Points
    to_label = "प्रति :" if lang == "hi" else "To :"
    p = doc.add_paragraph(to_label)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    for designation in data.get("to_designations", []):
        p = doc.add_paragraph(designation, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    # Attached section
    attached_label = "संलग्न :" if lang == "hi" else "Attached :"
    p = doc.add_paragraph(f"{attached_label} {data.get('attached_pdf_name', '')}")
    p.runs[0].font.size = Pt(12)

    # Try to convert and add PDF pages as images
    uploaded_pdf_path = data.get("uploaded_pdf_path")
    if uploaded_pdf_path and os.path.exists(uploaded_pdf_path):
        try:
            from pdf2image import convert_from_path

            temp_dir = os.path.join(settings.MEDIA_ROOT, "policy_uploads")
            os.makedirs(temp_dir, exist_ok=True)

            doc.add_page_break()

            try:
                images = convert_from_path(uploaded_pdf_path, dpi=200, fmt='png')

                for idx, image in enumerate(images):
                    temp_img_path = os.path.join(temp_dir, f"temp_page_{uuid.uuid4().hex}_{idx}.png")

                    try:
                        image.save(temp_img_path, "PNG")
                        if idx > 0:
                            doc.add_page_break()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(temp_img_path, width=Inches(6.5))
                        try:
                            os.remove(temp_img_path)
                        except OSError as rm_error:
                            logger.warning("Cleanup error: %s", rm_error)
                    except Exception as img_error:
                        logger.warning("Image processing error for page %d: %s", idx, img_error)
                        continue

                try:
                    os.remove(uploaded_pdf_path)
                except OSError as cleanup_error:
                    logger.warning("PDF cleanup error: %s", cleanup_error)

            except Exception as conv_error:
                logger.warning("PDF conversion error (poppler likely missing): %s", conv_error)
                doc.add_page_break()
                p = doc.add_paragraph("Note: Attached PDF pages could not be converted to images.")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p = doc.add_paragraph("This requires 'poppler' to be installed on your system.")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(10)
                p = doc.add_paragraph("Please download the complete document as PDF to view all pages.")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(10)

        except ImportError as imp_err:
            logger.warning("pdf2image not available: %s", imp_err)
            doc.add_page_break()
            p = doc.add_paragraph("[PDF attachment module not available - download as PDF to view complete document]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
        except Exception as exc:
            logger.exception("Error processing PDF attachment: %s", exc)
            doc.add_page_break()
            p = doc.add_paragraph("[Error loading attached PDF - download as PDF for complete document]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="Policy.docx"'
    return response
